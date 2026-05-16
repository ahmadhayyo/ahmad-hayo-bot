"""
Testing & validation tools — validate documents, run scripts, open & screenshot files.

Gives the agent the ability to test files and applications sent by the user:
  - validate_document: deep integrity check for Excel/Word/PDF/text files
  - run_executable_test: run a script or executable and capture output
  - open_and_screenshot: open a file with its default app and take a screenshot
"""

from __future__ import annotations

import datetime
import logging
import os
import subprocess
import time
from pathlib import Path
from typing import Annotated

from langchain_core.tools import tool

from config import DESKTOP_DIR, PS_TIMEOUT

logger = logging.getLogger("hayo.tools.testing")


@tool
def validate_document(path: str) -> str:
    """فحص شامل لملف والتحقق من سلامته ومحتوياته.

    يفحص الملف بعمق: هل هو سليم؟ هل يمكن فتحه؟ ما محتوياته؟
    يعمل مع: Excel (.xlsx), Word (.docx), PDF (.pdf), نصوص، صور، وغيرها.
    يعرض تقرير مفصل عن حالة الملف ومحتوياته.

    Args:
        path: مسار الملف المراد فحصه
    """
    p = Path(os.path.expandvars(os.path.expanduser(path))).resolve()
    if not p.exists():
        return f"❌ FAIL: file not found: {path}"

    stat = p.stat()
    size_kb = stat.st_size / 1024
    modified = datetime.datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
    suffix = p.suffix.lower()

    report = [
        "═══ تقرير فحص الملف ═══",
        f"📄 الملف: {p.name}",
        f"📂 المسار: {p}",
        f"📏 الحجم: {size_kb:.1f} KB ({stat.st_size:,} bytes)",
        f"📅 آخر تعديل: {modified}",
        f"🏷️ النوع: {suffix or 'unknown'}",
        "",
    ]

    if stat.st_size == 0:
        report.append("❌ FAIL: الملف فارغ (0 bytes)")
        return "\n".join(report)

    checks_passed = 0
    checks_total = 0

    # Excel validation
    if suffix in (".xlsx", ".xls"):
        checks_total += 3
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(p), read_only=True, data_only=True)
            report.append("✅ CHECK 1/3: الملف يُفتح بنجاح")
            checks_passed += 1

            if wb.sheetnames:
                report.append(f"✅ CHECK 2/3: يحتوي على {len(wb.sheetnames)} ورقة: {', '.join(wb.sheetnames)}")
                checks_passed += 1
            else:
                report.append("⚠️ CHECK 2/3: لا توجد أوراق عمل")

            total_rows = 0
            total_data_cells = 0
            for ws in wb.worksheets:
                rows = ws.max_row or 0
                cols = ws.max_column or 0
                total_rows += rows
                report.append(f"   📊 {ws.title}: {rows} صف × {cols} عمود")
                for row in ws.iter_rows(max_row=min(rows, 5)):
                    for cell in row:
                        if cell.value is not None:
                            total_data_cells += 1

            if total_data_cells > 0:
                report.append(f"✅ CHECK 3/3: يحتوي على بيانات ({total_data_cells}+ خلية بها قيم)")
                checks_passed += 1
            else:
                report.append("⚠️ CHECK 3/3: لا توجد بيانات في أول 5 صفوف")

            wb.close()
        except Exception as exc:
            report.append(f"❌ FAIL: لا يمكن فتح الملف: {type(exc).__name__}: {exc}")

    # Word validation
    elif suffix == ".docx":
        checks_total += 3
        try:
            from docx import Document
            doc = Document(str(p))
            report.append("✅ CHECK 1/3: الملف يُفتح بنجاح")
            checks_passed += 1

            para_count = len(doc.paragraphs)
            non_empty = [p for p in doc.paragraphs if p.text.strip()]
            report.append(f"✅ CHECK 2/3: {para_count} فقرة ({len(non_empty)} غير فارغة)")
            checks_passed += 1

            table_count = len(doc.tables)
            word_count = sum(len(para.text.split()) for para in doc.paragraphs)
            has_content = word_count > 0 or table_count > 0

            details = []
            if word_count > 0:
                details.append(f"{word_count} كلمة")
            if table_count > 0:
                details.append(f"{table_count} جدول")
                for i, table in enumerate(doc.tables):
                    rows = len(table.rows)
                    cols = len(table.columns)
                    report.append(f"   📊 جدول {i+1}: {rows} صف × {cols} عمود")

            # Check for images
            try:
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                image_count = 0
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        image_count += 1
                if image_count > 0:
                    details.append(f"{image_count} صورة")
            except Exception:
                pass

            if has_content:
                report.append(f"✅ CHECK 3/3: يحتوي على محتوى ({', '.join(details)})")
                checks_passed += 1
            else:
                report.append("⚠️ CHECK 3/3: المستند فارغ من المحتوى")

            # Preview first 3 paragraphs
            preview = [para.text[:100] for para in non_empty[:3]]
            if preview:
                report.append("\n📝 معاينة:")
                for i, text in enumerate(preview, 1):
                    report.append(f"   {i}. {text}")
        except Exception as exc:
            report.append(f"❌ FAIL: لا يمكن فتح الملف: {type(exc).__name__}: {exc}")

    # PDF validation
    elif suffix == ".pdf":
        checks_total += 3
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(p))
            report.append("✅ CHECK 1/3: الملف يُفتح بنجاح")
            checks_passed += 1

            page_count = len(reader.pages)
            if page_count > 0:
                report.append(f"✅ CHECK 2/3: يحتوي على {page_count} صفحة")
                checks_passed += 1
            else:
                report.append("⚠️ CHECK 2/3: لا توجد صفحات")

            first_text = reader.pages[0].extract_text() if reader.pages else ""
            if first_text and first_text.strip():
                report.append(f"✅ CHECK 3/3: يحتوي على نص قابل للاستخراج")
                checks_passed += 1
                report.append(f"\n📝 معاينة الصفحة الأولى:\n   {first_text[:300]}")
            else:
                report.append("⚠️ CHECK 3/3: لا يوجد نص قابل للاستخراج (قد يكون صوراً)")
        except Exception as exc:
            report.append(f"❌ FAIL: لا يمكن فتح الملف: {type(exc).__name__}: {exc}")

    # Image validation
    elif suffix in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".svg"):
        checks_total += 2
        try:
            from PIL import Image
            img = Image.open(str(p))
            report.append(f"✅ CHECK 1/2: الصورة تُفتح بنجاح")
            checks_passed += 1
            report.append(f"✅ CHECK 2/2: الأبعاد {img.size[0]}×{img.size[1]}, النمط: {img.mode}")
            checks_passed += 1
            img.close()
        except ImportError:
            report.append("⚠️ مكتبة Pillow غير متوفرة لفحص الصور")
        except Exception as exc:
            report.append(f"❌ FAIL: لا يمكن فتح الصورة: {type(exc).__name__}: {exc}")

    # Text/code file validation
    elif suffix in (".txt", ".csv", ".json", ".xml", ".html", ".md", ".py", ".js", ".ts", ".css", ".yaml", ".yml"):
        checks_total += 2
        try:
            with open(str(p), "r", encoding="utf-8", errors="replace") as f:
                content = f.read(10000)
            line_count = content.count("\n") + 1
            char_count = len(content)
            report.append(f"✅ CHECK 1/2: الملف يُقرأ بنجاح ({line_count} سطر)")
            checks_passed += 1

            if char_count > 0:
                report.append(f"✅ CHECK 2/2: يحتوي على {char_count:,} حرف")
                checks_passed += 1
            else:
                report.append("⚠️ CHECK 2/2: الملف فارغ")

            # JSON validation
            if suffix == ".json":
                import json
                try:
                    json.loads(content)
                    report.append("✅ BONUS: JSON صالح")
                except json.JSONDecodeError as je:
                    report.append(f"⚠️ BONUS: JSON غير صالح: {je}")

            report.append(f"\n📝 معاينة:\n{content[:500]}")
        except Exception as exc:
            report.append(f"❌ FAIL: لا يمكن قراءة الملف: {type(exc).__name__}: {exc}")

    # Unknown file type
    else:
        checks_total += 1
        report.append(f"ℹ️ نوع الملف ({suffix}) غير معروف — فحص أساسي فقط")
        if stat.st_size > 0:
            report.append("✅ CHECK 1/1: الملف موجود وغير فارغ")
            checks_passed += 1
        else:
            report.append("❌ CHECK 1/1: الملف فارغ")

    # Summary
    report.append("")
    report.append("═══ النتيجة ═══")
    if checks_total > 0:
        pct = (checks_passed / checks_total) * 100
        status = "✅ ناجح" if checks_passed == checks_total else "⚠️ ناجح جزئياً" if checks_passed > 0 else "❌ فاشل"
        report.append(f"{status}: {checks_passed}/{checks_total} فحوصات ({pct:.0f}%)")
    else:
        report.append("ℹ️ لم يتم إجراء فحوصات")

    return "\n".join(report)


@tool
def run_executable_test(
    command: str,
    workdir: Annotated[str, "مجلد العمل. استخدم '.' للمجلد الحالي أو 'desktop:' لسطح المكتب."] = ".",
    timeout: Annotated[int, "الحد الأقصى للتنفيذ بالثواني (افتراضي: 30)."] = 30,
) -> str:
    """تشغيل ملف تنفيذي أو سكريبت واختبار نتيجته.

    يشغّل الأمر ويلتقط المخرجات (stdout + stderr) ورمز الخروج.
    مناسب لاختبار: سكريبتات Python، ملفات batch، برامج، أوامر npm/pip، وغيرها.

    Args:
        command: الأمر المراد تشغيله (مثل 'python script.py' أو 'node app.js')
        workdir: مجلد العمل
        timeout: الحد الأقصى بالثواني
    """
    # Resolve working directory
    w = workdir.strip()
    if not w or w == ".":
        resolved_dir = str(DESKTOP_DIR)
    elif w.lower() in ("desktop", "desktop:"):
        resolved_dir = str(DESKTOP_DIR)
    else:
        resolved_dir = str(Path(os.path.expandvars(os.path.expanduser(w))).resolve())

    if not os.path.isdir(resolved_dir):
        return f"❌ مجلد العمل غير موجود: {resolved_dir}"

    timeout = max(1, min(timeout, PS_TIMEOUT))

    report = [
        "═══ تقرير اختبار التنفيذ ═══",
        f"🔧 الأمر: {command}",
        f"📂 المجلد: {resolved_dir}",
        f"⏱️ الحد الأقصى: {timeout}s",
        "",
    ]

    start_time = time.time()
    try:
        completed = subprocess.run(
            command,
            cwd=resolved_dir,
            capture_output=True,
            text=True,
            timeout=timeout,
            shell=True,
        )
        elapsed = time.time() - start_time

        report.append(f"⏱️ مدة التنفيذ: {elapsed:.2f}s")
        report.append(f"📊 رمز الخروج: {completed.returncode}")

        if completed.returncode == 0:
            report.append("✅ النتيجة: نجح التنفيذ")
        else:
            report.append(f"❌ النتيجة: فشل التنفيذ (exit code {completed.returncode})")

        if completed.stdout:
            stdout = completed.stdout[:4000]
            report.append(f"\n📤 المخرجات (stdout):\n{stdout}")
            if len(completed.stdout) > 4000:
                report.append(f"... ({len(completed.stdout) - 4000} حرف إضافي محذوف)")

        if completed.stderr:
            stderr = completed.stderr[:2000]
            report.append(f"\n⚠️ الأخطاء (stderr):\n{stderr}")
            if len(completed.stderr) > 2000:
                report.append(f"... ({len(completed.stderr) - 2000} حرف إضافي محذوف)")

    except subprocess.TimeoutExpired:
        elapsed = time.time() - start_time
        report.append(f"⏱️ مدة التنفيذ: {elapsed:.2f}s")
        report.append(f"❌ النتيجة: انتهت المهلة ({timeout}s)")
    except Exception as exc:
        report.append(f"❌ خطأ في التنفيذ: {type(exc).__name__}: {exc}")

    return "\n".join(report)


@tool
def open_and_screenshot(
    path: str,
    wait_seconds: Annotated[int, "ثواني الانتظار بعد الفتح قبل التقاط الشاشة (افتراضي: 3)."] = 3,
    screenshot_name: Annotated[str, "اسم ملف لقطة الشاشة (افتراضي: test_screenshot.png)."] = "",
) -> str:
    """فتح ملف بالتطبيق الافتراضي وأخذ لقطة شاشة للتحقق البصري.

    يفتح الملف باستخدام التطبيق الافتراضي في Windows، ينتظر ليتحميل،
    ثم يأخذ لقطة شاشة ويحفظها على سطح المكتب.
    مناسب للتحقق البصري من: مستندات، صور، تطبيقات، مواقع HTML.

    Args:
        path: مسار الملف المراد فتحه
        wait_seconds: ثواني الانتظار قبل التقاط الشاشة
        screenshot_name: اسم ملف لقطة الشاشة
    """
    p = Path(os.path.expandvars(os.path.expanduser(path))).resolve()
    if not p.exists():
        return f"❌ الملف غير موجود: {path}"

    wait_seconds = max(1, min(wait_seconds, 30))

    report = [
        "═══ تقرير الفتح والتقاط الشاشة ═══",
        f"📄 الملف: {p.name}",
        f"📂 المسار: {p}",
        "",
    ]

    # Open the file with default application
    try:
        subprocess.Popen(
            ["cmd.exe", "/C", "start", "", str(p)],
            shell=False,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        report.append(f"✅ تم فتح الملف بالتطبيق الافتراضي")
    except Exception as exc:
        report.append(f"❌ فشل فتح الملف: {type(exc).__name__}: {exc}")
        return "\n".join(report)

    # Wait for the application to load
    report.append(f"⏳ انتظار {wait_seconds} ثانية للتحميل...")
    time.sleep(wait_seconds)

    # Take screenshot
    screenshot_file = screenshot_name or f"test_{p.stem}.png"
    if not screenshot_file.endswith(".png"):
        screenshot_file += ".png"
    screenshot_path = DESKTOP_DIR / screenshot_file

    try:
        import pyautogui
        img = pyautogui.screenshot()
        screenshot_path.parent.mkdir(parents=True, exist_ok=True)
        img.save(str(screenshot_path))
        report.append(f"📸 تم حفظ لقطة الشاشة: {screenshot_path}")
        report.append(f"   الأبعاد: {img.size[0]}×{img.size[1]}")
    except ImportError:
        report.append("⚠️ مكتبة pyautogui غير متوفرة — لا يمكن التقاط الشاشة")
    except Exception as exc:
        report.append(f"❌ فشل التقاط الشاشة: {type(exc).__name__}: {exc}")

    report.append("")
    report.append("✅ اكتمل الاختبار البصري — راجع لقطة الشاشة على سطح المكتب")

    return "\n".join(report)
