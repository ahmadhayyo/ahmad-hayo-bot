"""
Office document tools — Excel, Word, PDF, Translation.

Create, read, edit spreadsheets, documents, and PDF files.
Uses openpyxl (Excel), python-docx (Word), pypdf + reportlab (PDF).
Includes translation tools using deep-translator (Google Translate).
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from langchain_core.tools import tool


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL TOOLS
# ═══════════════════════════════════════════════════════════════════════════════

@tool
def excel_create(path: str, data: str, sheet_name: str = "Sheet1") -> str:
    """إنشاء ملف Excel جديد من بيانات JSON.

    Args:
        path: مسار الملف (مثل C:/Users/user/Desktop/report.xlsx)
        data: بيانات JSON — قائمة من القوائم أو قائمة من القواميس.
              مثال 1: [["الاسم","الراتب"],["أحمد",5000],["سارة",6000]]
              مثال 2: [{"الاسم":"أحمد","الراتب":5000},{"الاسم":"سارة","الراتب":6000}]
        sheet_name: اسم الورقة (افتراضي: Sheet1)
    """
    try:
        from openpyxl import Workbook
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    try:
        rows = json.loads(data) if isinstance(data, str) else data
    except json.JSONDecodeError:
        return f"Error: invalid JSON data"

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    if rows and isinstance(rows[0], dict):
        headers = list(rows[0].keys())
        ws.append(headers)
        for row in rows:
            ws.append([row.get(h, "") for h in headers])
    elif rows and isinstance(rows[0], (list, tuple)):
        for row in rows:
            ws.append(list(row))
    else:
        return "Error: data must be list of lists or list of dicts"

    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(p))
    return f"Excel file created: {p} ({len(rows)} rows)"


@tool
def excel_read(path: str, sheet_name: str = "") -> str:
    """قراءة ملف Excel وإرجاع البيانات كنص.

    Args:
        path: مسار ملف Excel
        sheet_name: اسم الورقة (فارغ = الورقة النشطة)
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    p = Path(path)
    if not p.exists():
        return f"Error: file not found: {path}"

    wb = load_workbook(str(p), data_only=True)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append([str(c) if c is not None else "" for c in row])

    if not rows:
        return "Empty spreadsheet"

    result = f"Sheet: {ws.title} | Rows: {len(rows)} | Columns: {len(rows[0])}\n"
    result += f"Sheets available: {wb.sheetnames}\n\n"

    col_widths = [0] * len(rows[0])
    for row in rows[:50]:
        for i, c in enumerate(row):
            col_widths[i] = max(col_widths[i], len(c))

    for i, row in enumerate(rows[:100]):
        line = " | ".join(c.ljust(min(w, 30)) for c, w in zip(row, col_widths))
        result += line + "\n"
        if i == 0:
            result += "-" * len(line) + "\n"

    if len(rows) > 100:
        result += f"\n... ({len(rows) - 100} more rows)"
    return result


@tool
def excel_edit(path: str, cell: str, value: str, sheet_name: str = "") -> str:
    """تعديل خلية في ملف Excel.

    Args:
        path: مسار ملف Excel
        cell: عنوان الخلية (مثل A1, B5, C10)
        value: القيمة الجديدة
        sheet_name: اسم الورقة (فارغ = الورقة النشطة)
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    p = Path(path)
    if not p.exists():
        return f"Error: file not found: {path}"

    wb = load_workbook(str(p))
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    old_val = ws[cell].value
    try:
        ws[cell] = float(value)
    except ValueError:
        ws[cell] = value

    wb.save(str(p))
    return f"Cell {cell} updated: '{old_val}' → '{value}'"


@tool
def excel_add_rows(path: str, data: str, sheet_name: str = "") -> str:
    """إضافة صفوف جديدة إلى ملف Excel موجود.

    Args:
        path: مسار ملف Excel
        data: بيانات JSON — قائمة من القوائم أو قائمة من القواميس
        sheet_name: اسم الورقة (فارغ = الورقة النشطة)
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    p = Path(path)
    if not p.exists():
        return f"Error: file not found: {path}"

    try:
        rows = json.loads(data) if isinstance(data, str) else data
    except json.JSONDecodeError:
        return "Error: invalid JSON data"

    wb = load_workbook(str(p))
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    count = 0
    if rows and isinstance(rows[0], dict):
        headers = [c.value for c in ws[1]] if ws.max_row >= 1 else list(rows[0].keys())
        for row in rows:
            ws.append([row.get(str(h), "") for h in headers])
            count += 1
    else:
        for row in rows:
            ws.append(list(row))
            count += 1

    wb.save(str(p))
    return f"Added {count} rows to {p}"


@tool
def excel_add_column(path: str, header: str, formula_or_values: str, sheet_name: str = "") -> str:
    """إضافة عمود جديد إلى ملف Excel.

    Args:
        path: مسار ملف Excel
        header: اسم العمود الجديد
        formula_or_values: صيغة Excel (تبدأ بـ =) أو JSON قائمة قيم.
                          مثال صيغة: "=B{row}*0.1" (يتم استبدال {row} برقم الصف)
                          مثال قيم: "[500, 600, 450]"
        sheet_name: اسم الورقة
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    p = Path(path)
    if not p.exists():
        return f"Error: file not found: {path}"

    wb = load_workbook(str(p))
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    new_col = ws.max_column + 1
    ws.cell(row=1, column=new_col, value=header)

    if formula_or_values.strip().startswith("="):
        for row_num in range(2, ws.max_row + 1):
            formula = formula_or_values.replace("{row}", str(row_num))
            ws.cell(row=row_num, column=new_col, value=formula)
    else:
        try:
            values = json.loads(formula_or_values)
            for i, val in enumerate(values, start=2):
                ws.cell(row=i, column=new_col, value=val)
        except json.JSONDecodeError:
            for row_num in range(2, ws.max_row + 1):
                ws.cell(row=row_num, column=new_col, value=formula_or_values)

    wb.save(str(p))
    return f"Column '{header}' added at position {new_col}"


# ═══════════════════════════════════════════════════════════════════════════════
#  WORD TOOLS
# ═══════════════════════════════════════════════════════════════════════════════

@tool
def word_create(path: str, content: str, title: str = "") -> str:
    """إنشاء ملف Word جديد.

    Args:
        path: مسار الملف (مثل C:/Users/user/Desktop/report.docx)
        content: المحتوى — نص عادي أو Markdown بسيط.
                 استخدم # للعناوين، - للقوائم، سطر فارغ لفقرة جديدة.
        title: عنوان المستند (اختياري)
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped[0:3].rstrip(". ").isdigit() and ". " in stripped[:5]:
            idx = stripped.index(". ")
            doc.add_paragraph(stripped[idx + 2:], style="List Number")
        else:
            doc.add_paragraph(stripped)

    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(p))
    return f"Word file created: {p}"


@tool
def word_read(path: str) -> str:
    """قراءة محتوى ملف Word.

    Args:
        path: مسار ملف Word (.docx)
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    p = Path(path)
    if not p.exists():
        return f"Error: file not found: {path}"

    doc = Document(str(p))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style else ""
        if "heading 1" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        elif "list bullet" in style:
            lines.append(f"- {text}")
        elif "list number" in style:
            lines.append(f"1. {text}")
        else:
            lines.append(text)

    # Also read tables
    for i, table in enumerate(doc.tables):
        lines.append(f"\n[Table {i + 1}]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(cells))

    return "\n".join(lines) if lines else "Empty document"


@tool
def word_edit(path: str, find_text: str, replace_text: str) -> str:
    """البحث والاستبدال في ملف Word.

    Args:
        path: مسار ملف Word
        find_text: النص المراد البحث عنه
        replace_text: النص البديل
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    p = Path(path)
    if not p.exists():
        return f"Error: file not found: {path}"

    doc = Document(str(p))
    count = 0
    for para in doc.paragraphs:
        if find_text in para.text:
            for run in para.runs:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    count += 1

    doc.save(str(p))
    return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'"


# ═══════════════════════════════════════════════════════════════════════════════
#  PDF TOOLS
# ═══════════════════════════════════════════════════════════════════════════════

@tool
def pdf_read(path: str, max_pages: int = 50) -> str:
    """قراءة محتوى ملف PDF واستخراج النص.

    Args:
        path: مسار ملف PDF
        max_pages: الحد الأقصى لعدد الصفحات المقروءة (افتراضي 50)
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return "Error: pypdf not installed. Run: pip install pypdf"

    p = Path(path)
    if not p.exists():
        return f"Error: file not found: {path}"

    reader = PdfReader(str(p))
    total = len(reader.pages)

    lines = [f"PDF: {p.name} | Pages: {total}\n"]
    for i, page in enumerate(reader.pages[:max_pages]):
        text = page.extract_text() or ""
        if text.strip():
            lines.append(f"--- Page {i + 1} ---")
            lines.append(text.strip())

    if total > max_pages:
        lines.append(f"\n... ({total - max_pages} more pages)")

    return "\n".join(lines) if len(lines) > 1 else "No text found in PDF (may be scanned/image-based)"


@tool
def pdf_create(path: str, content: str, title: str = "") -> str:
    """إنشاء ملف PDF من نص.

    Args:
        path: مسار الملف (مثل C:/Users/user/Desktop/report.pdf)
        content: المحتوى النصي (كل سطر = فقرة)
        title: عنوان المستند (اختياري)
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return "Error: reportlab not installed. Run: pip install reportlab"

    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(p), pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    if title:
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 12))

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
        elif stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], styles["Heading1"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], styles["Heading2"]))
        else:
            story.append(Paragraph(stripped, styles["Normal"]))

    doc.build(story)
    return f"PDF file created: {p}"


@tool
def pdf_merge(paths: str, output_path: str) -> str:
    """دمج عدة ملفات PDF في ملف واحد.

    Args:
        paths: قائمة مسارات PDF مفصولة بفاصلة أو JSON list
        output_path: مسار الملف الناتج
    """
    try:
        from pypdf import PdfWriter
    except ImportError:
        return "Error: pypdf not installed. Run: pip install pypdf"

    try:
        file_list = json.loads(paths)
    except (json.JSONDecodeError, TypeError):
        file_list = [p.strip() for p in paths.split(",")]

    writer = PdfWriter()
    for fp in file_list:
        fp = fp.strip()
        if not Path(fp).exists():
            return f"Error: file not found: {fp}"
        writer.append(fp)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    writer.write(str(out))
    return f"Merged {len(file_list)} PDFs into: {out}"


@tool
def convert_excel_to_pdf(excel_path: str, pdf_path: str) -> str:
    """تحويل ملف Excel إلى PDF.

    Args:
        excel_path: مسار ملف Excel المصدر
        pdf_path: مسار ملف PDF الناتج
    """
    try:
        from openpyxl import load_workbook
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
        from reportlab.lib import colors
    except ImportError:
        return "Error: openpyxl and reportlab required. Run: pip install openpyxl reportlab"

    p = Path(excel_path)
    if not p.exists():
        return f"Error: file not found: {excel_path}"

    wb = load_workbook(str(p), data_only=True)
    ws = wb.active

    data = []
    for row in ws.iter_rows(values_only=True):
        data.append([str(c) if c is not None else "" for c in row])

    if not data:
        return "Error: spreadsheet is empty"

    out = Path(pdf_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(out), pagesize=landscape(A4))
    table = Table(data)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#D9E2F3")]),
    ]))

    doc.build([table])
    return f"Converted {excel_path} → {pdf_path} ({len(data)} rows)"


@tool
def convert_word_to_pdf(word_path: str, pdf_path: str) -> str:
    """تحويل ملف Word إلى PDF.

    Args:
        word_path: مسار ملف Word المصدر
        pdf_path: مسار ملف PDF الناتج
    """
    try:
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        return "Error: python-docx and reportlab required. Run: pip install python-docx reportlab"

    p = Path(word_path)
    if not p.exists():
        return f"Error: file not found: {word_path}"

    doc = Document(str(p))
    styles = getSampleStyleSheet()
    story = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue
        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name:
            story.append(Paragraph(text, styles["Heading1"]))
        elif "heading 2" in style_name:
            story.append(Paragraph(text, styles["Heading2"]))
        else:
            story.append(Paragraph(text, styles["Normal"]))

    out = Path(pdf_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    pdf_doc = SimpleDocTemplate(str(out), pagesize=A4)
    pdf_doc.build(story)
    return f"Converted {word_path} → {pdf_path}"


# ═══════════════════════════════════════════════════════════════════════════════
#  TRANSLATION TOOLS
# ═══════════════════════════════════════════════════════════════════════════════

# Supported language codes for quick reference in tool docstrings
_LANG_EXAMPLES = (
    "ar=عربي, hi=هندي, en=إنجليزي, fr=فرنسي, es=إسباني, de=ألماني, "
    "tr=تركي, fa=فارسي, ur=أردو, zh-CN=صيني, ja=ياباني, ko=كوري, "
    "ru=روسي, pt=برتغالي, it=إيطالي, nl=هولندي, auto=كشف تلقائي"
)


def _translate_single(text: str, source: str, target: str) -> str:
    """Translate a single string using deep-translator GoogleTranslator."""
    from deep_translator import GoogleTranslator
    if not text or not text.strip():
        return text
    try:
        return GoogleTranslator(source=source, target=target).translate(text)
    except Exception:
        return text


def _translate_batch(texts: list[str], source: str, target: str) -> list[str]:
    """Translate a list of strings, preserving empty/whitespace entries."""
    from deep_translator import GoogleTranslator
    translator = GoogleTranslator(source=source, target=target)
    results: list[str] = []
    # deep-translator supports batch up to ~5000 chars; chunk manually
    batch: list[tuple[int, str]] = []
    for i, t in enumerate(texts):
        if t and t.strip():
            batch.append((i, t))

    translated_map: dict[int, str] = {}
    # Translate in chunks of 50 to avoid API limits
    chunk_size = 50
    for start in range(0, len(batch), chunk_size):
        chunk = batch[start:start + chunk_size]
        chunk_texts = [t for _, t in chunk]
        try:
            translated = translator.translate_batch(chunk_texts)
            for (idx, _orig), tr_text in zip(chunk, translated):
                translated_map[idx] = tr_text if tr_text else _orig
        except Exception:
            # Fallback: translate one by one
            for idx, orig in chunk:
                try:
                    translated_map[idx] = translator.translate(orig)
                except Exception:
                    translated_map[idx] = orig

    for i, t in enumerate(texts):
        results.append(translated_map.get(i, t))
    return results


@tool
def translate_text(text: str, target_lang: str, source_lang: str = "auto") -> str:
    """ترجمة نص من لغة إلى أخرى باستخدام Google Translate.

    Args:
        text: النص المراد ترجمته
        target_lang: رمز اللغة الهدف (مثل hi=هندي, en=إنجليزي, ar=عربي, fr=فرنسي)
        source_lang: رمز اللغة المصدر (افتراضي: auto = كشف تلقائي)

    أمثلة:
        translate_text(text='مرحبا بالعالم', target_lang='hi')
        translate_text(text='Hello World', target_lang='ar')
        translate_text(text='مرحبا', target_lang='en', source_lang='ar')

    اللغات المدعومة: ar, hi, en, fr, es, de, tr, fa, ur, zh-CN, ja, ko, ru, pt, it, nl وغيرها
    """
    try:
        from deep_translator import GoogleTranslator
    except ImportError:
        return "Error: deep-translator not installed. Run: pip install deep-translator"

    if not text or not text.strip():
        return "Error: empty text provided"

    try:
        result = GoogleTranslator(source=source_lang, target=target_lang).translate(text)
        return result if result else "Error: translation returned empty result"
    except Exception as exc:
        return f"Error: {type(exc).__name__}: {exc}"


@tool
def excel_clone_translated(
    source_path: str,
    dest_path: str,
    target_lang: str,
    source_lang: str = "auto",
    sheets: str = "",
) -> str:
    """استنساخ ملف Excel مع ترجمة المحتوى النصي والحفاظ على التنسيق الكامل.

    يقرأ ملف Excel المصدر، يترجم جميع الخلايا النصية، ويحفظ نسخة جديدة
    مع الحفاظ الكامل على: الألوان، الخطوط، الحدود، عرض الأعمدة، ارتفاع الصفوف،
    الخلايا المدمجة، التنسيق الشرطي، والصيغ الرياضية.

    Args:
        source_path: مسار ملف Excel المصدر
        dest_path: مسار ملف Excel الناتج المترجم
        target_lang: رمز اللغة الهدف (مثل hi=هندي, en=إنجليزي, fr=فرنسي)
        source_lang: رمز اللغة المصدر (افتراضي: auto = كشف تلقائي)
        sheets: أسماء الأوراق للترجمة مفصولة بفاصلة (فارغ = كل الأوراق)

    أمثلة:
        excel_clone_translated(
            source_path='C:/Users/user/Desktop/report.xlsx',
            dest_path='C:/Users/user/Desktop/report_hindi.xlsx',
            target_lang='hi', source_lang='ar'
        )
        excel_clone_translated(
            source_path='~/Desktop/data.xlsx',
            dest_path='~/Desktop/data_english.xlsx',
            target_lang='en'
        )

    اللغات المدعومة: ar, hi, en, fr, es, de, tr, fa, ur, zh-CN, ja, ko, ru, pt, it, nl وغيرها
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    try:
        from deep_translator import GoogleTranslator
    except ImportError:
        return "Error: deep-translator not installed. Run: pip install deep-translator"

    src = Path(os.path.expandvars(os.path.expanduser(source_path))).resolve()
    if not src.exists():
        return f"Error: file not found: {source_path}"

    dst = Path(os.path.expandvars(os.path.expanduser(dest_path))).resolve()
    dst.parent.mkdir(parents=True, exist_ok=True)

    try:
        wb = load_workbook(str(src))
    except Exception as exc:
        return f"Error loading workbook: {type(exc).__name__}: {exc}"

    # Determine which sheets to translate
    if sheets:
        sheet_names = [s.strip() for s in sheets.split(",") if s.strip()]
    else:
        sheet_names = wb.sheetnames

    translator = GoogleTranslator(source=source_lang, target=target_lang)
    total_translated = 0
    total_cells = 0

    for sheet_name in sheet_names:
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]

        # Collect all text cells for batch translation
        text_cells: list[tuple[int, int, str]] = []  # (row, col, value)
        for row in ws.iter_rows():
            for cell in row:
                total_cells += 1
                if cell.value is not None and isinstance(cell.value, str):
                    val = cell.value.strip()
                    # Skip formulas, empty strings, pure numbers
                    if val and not val.startswith("="):
                        text_cells.append((cell.row, cell.column, val))

        if not text_cells:
            continue

        # Batch translate for efficiency
        originals = [t[2] for t in text_cells]
        translated = _translate_batch(originals, source_lang, target_lang)

        # Write translated values back (formatting is preserved automatically)
        for (row, col, _orig), tr_text in zip(text_cells, translated):
            ws.cell(row=row, column=col).value = tr_text
            total_translated += 1

    # Also translate the sheet tab names
    for ws in wb.worksheets:
        if ws.title in sheet_names:
            original_title = ws.title
            try:
                new_title = translator.translate(original_title)
                if new_title and new_title != original_title:
                    ws.title = new_title[:31]  # Excel sheet name max 31 chars
            except Exception:
                pass

    try:
        wb.save(str(dst))
    except Exception as exc:
        return f"Error saving workbook: {type(exc).__name__}: {exc}"

    return (
        f"Excel file cloned and translated successfully!\n"
        f"Source: {src}\n"
        f"Output: {dst}\n"
        f"Language: {source_lang} → {target_lang}\n"
        f"Sheets translated: {len(sheet_names)}\n"
        f"Cells translated: {total_translated} / {total_cells} total"
    )
